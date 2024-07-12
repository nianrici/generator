from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def copiar_contenido(tabla):
    contenido = {}
    for fila in tabla.rows:
        etiqueta = fila.cells[0].text.strip()
        contenido[etiqueta] = fila.cells[1]
    return contenido

def copiar_tabla(tabla_origen):
    # Crear una nueva tabla
    tbl = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    tbl.append(tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), str(Inches(6).twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)

    # Copiar filas y celdas
    for fila_origen in tabla_origen.rows:
        tr = OxmlElement('w:tr')
        tbl.append(tr)
        for celda_origen in fila_origen.cells:
            tc = OxmlElement('w:tc')
            tr.append(tc)
            tcPr = OxmlElement('w:tcPr')
            tc.append(tcPr)
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), '0')
            tcW.set(qn('w:type'), 'auto')
            tcPr.append(tcW)
            p = OxmlElement('w:p')
            tc.append(p)
            r = OxmlElement('w:r')
            p.append(r)
            t = OxmlElement('w:t')
            r.append(t)
            t.text = celda_origen.text

    return tbl

def reemplazar_etiquetas(documento, contenido):
    for parrafo in documento.paragraphs:
        for etiqueta, celda in contenido.items():
            if etiqueta in parrafo.text:
                if celda.tables:
                    # Si hay una tabla, la copiamos
                    elemento_parrafo = parrafo._element
                    elemento_parrafo.addnext(OxmlElement('w:p'))  # Salto de línea antes
                    for tabla in celda.tables:
                        elemento_tabla = copiar_tabla(tabla)
                        elemento_parrafo.addnext(elemento_tabla)
                        elemento_parrafo = elemento_tabla
                    elemento_parrafo.addnext(OxmlElement('w:p'))  # Salto de línea después
                    parrafo.text = parrafo.text.replace(etiqueta, '')
                elif celda.text:
                    parrafo.text = parrafo.text.replace(etiqueta, celda.text)
                else:
                    for run in celda.paragraphs[0].runs:
                        if run.element.find('.//w:drawing') is not None:
                            nueva_run = parrafo.add_run()
                            nueva_run._element.append(run.element.find('.//w:drawing'))
                    parrafo.text = parrafo.text.replace(etiqueta, '')

def main():
    # Abrir el primer archivo
    doc1 = Document('test1.docx')
    tabla = doc1.tables[0]
    contenido = copiar_contenido(tabla)

    # Abrir el segundo archivo
    doc2 = Document('test2.docx')
    reemplazar_etiquetas(doc2, contenido)

    # Guardar el resultado
    doc2.save('generado.docx')

if __name__ == '__main__':
    main()
